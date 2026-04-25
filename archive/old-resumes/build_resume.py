"""
Build Ena Dodski's resume as a branded .docx file.
Uses HCHC brand book colors, fonts (Cormorant Garamond + Jost),
and positioning from the Project Handoff document.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ESPRESSO = RGBColor(0x3C, 0x2A, 0x21)
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)  # Darker general text for readability
MOCHA = RGBColor(0xA3, 0x77, 0x62)
SAND = RGBColor(0xD4, 0xC5, 0xB2)
IVORY = RGBColor(0xF5, 0xF0, 0xEB)
GOLD = RGBColor(0xC4, 0xA2, 0x65)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Font names (must be installed on the system or embedded)
HEADING_FONT = "Cormorant Garamond SemiBold"
HEADING_FONT_ALT = "Cormorant Garamond"
BODY_FONT = "Jost"
BODY_FONT_ALT = "Calibri"

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.4)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

# White background for general viewing / LinkedIn / Indeed compatibility


def set_font(run, name=BODY_FONT, size=10, color=DARK_TEXT, bold=False, italic=False):
    """Helper to set run font properties."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Set east-asia and complex script font too
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def set_paragraph_shading(paragraph, hex_color):
    """Set paragraph background/shading."""
    pPr = paragraph._element.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    pPr.append(shd)


def add_thin_line():
    """Add a thin sand-colored horizontal line."""
    p = doc.add_paragraph()
    add_spacing(p, before=4, after=4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="D4C5B2"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_navy_bar():
    """Add a navy accent bar (thick line)."""
    p = doc.add_paragraph()
    add_spacing(p, before=0, after=6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ══════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ENA DODSKI")
set_font(run, name=HEADING_FONT, size=28, color=NAVY, bold=False)
run.font.letter_spacing = Pt(3)
add_spacing(p, before=0, after=0)

# Title line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DESIGN LEADERSHIP  ·  STRATEGIC OPERATIONS")
set_font(run, name=BODY_FONT, size=9, color=MOCHA)
run.font.letter_spacing = Pt(2)
add_spacing(p, before=0, after=1)

# Contact line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("San Antonio, TX   |   210-789-1051   |   ena.dodski@gmail.com")
set_font(run, name=BODY_FONT, size=8.5, color=DARK_TEXT)
add_spacing(p, before=0, after=0)

# Navy bar under header
add_navy_bar()


# ══════════════════════════════════════════════════════════
#  PROFILE SUMMARY
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Profile")
set_font(run, name=HEADING_FONT, size=14, color=NAVY)
add_spacing(p, before=4, after=1)
add_thin_line()

p = doc.add_paragraph()
run = p.add_run(
    "Design leader with over eight years of experience in high-volume production and semi-custom "
    "residential design across the Texas Hill Country market. I lead design centers, develop scalable "
    "systems, and drive revenue growth through strategic builder partnerships. My work positions homes "
    "that sell, scale, and feel elevated, bridging creative vision with operational execution."
)
set_font(run, size=11, color=DARK_TEXT)
add_spacing(p, before=2, after=4, line=15)


# ══════════════════════════════════════════════════════════
#  CORE COMPETENCIES
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Core Competencies")
set_font(run, name=HEADING_FONT, size=14, color=NAVY)
add_spacing(p, before=4, after=1)
add_thin_line()

# Competencies in a 3-column table
comp_table = doc.add_table(rows=1, cols=3)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove table borders
tbl = comp_table._element
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tbl.insert(0, tblPr)
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

competencies = [
    ("Strategic Leadership", "Procedure development\nOperational reporting\nTeam management\nPersonnel compliance"),
    ("Design Expertise", "New construction finishes\nMaterial specification\nTrend analysis\nFlooring, cabinetry, tile"),
    ("Business Development", "Revenue growth strategy\nBuilder & realtor relations\nClient retention\nMarket positioning"),
]

row = comp_table.rows[0]
for i, (title, items) in enumerate(competencies):
    cell = row.cells[i]
    # Clear default paragraph
    cell.paragraphs[0].clear()

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_font(run, name=BODY_FONT, size=9, color=NAVY, bold=True)
    run.font.letter_spacing = Pt(1)
    add_spacing(p, before=2, after=2)

    # Items
    for line in items.split("\n"):
        p = cell.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=9, color=ESPRESSO)
        add_spacing(p, before=0, after=0, line=13)


# ══════════════════════════════════════════════════════════
#  PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Professional Experience")
set_font(run, name=HEADING_FONT, size=14, color=NAVY)
add_spacing(p, before=6, after=1)
add_navy_bar()


def add_position(company, location, role, dates, bullets):
    """Add a position block."""
    # Company + location
    p = doc.add_paragraph()
    run = p.add_run(company)
    set_font(run, name=HEADING_FONT, size=11.5, color=ESPRESSO, bold=True)
    run = p.add_run(f"   {location}")
    set_font(run, size=9, color=MOCHA)
    add_spacing(p, before=8, after=0)

    # Role + dates
    p = doc.add_paragraph()
    run = p.add_run(role.upper())
    set_font(run, name=BODY_FONT, size=8.5, color=MOCHA)
    run.font.letter_spacing = Pt(1.5)
    run = p.add_run(f"      {dates}")
    set_font(run, size=8.5, color=MOCHA)
    add_spacing(p, before=0, after=3)

    # Bullets
    for bullet in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run("●  ")
        set_font(run, size=6, color=MOCHA)
        run = p.add_run(bullet)
        set_font(run, size=11, color=DARK_TEXT)
        add_spacing(p, before=1, after=1, line=14)


# United Finishes
add_position(
    "United Finishes (An ADG Company)", "San Antonio, TX",
    "Design Center Manager", "Nov 2024 – Present",
    [
        "Direct all design center operations for a portfolio of production and semi-custom projects ranging from $500K to $1.5M.",
        "Lead operational reporting, facility oversight, and the implementation of standardized design center procedures across the region.",
        "Drive business development initiatives to secure new builder partnerships and expand the company's regional market footprint.",
        "Manage high-level client consultations, aligning architectural plans with curated interior finishes that support builder timelines and buyer satisfaction.",
        "Develop and mentor design staff, overseeing personnel compliance and professional growth.",
    ]
)

# KB Home
add_position(
    "KB Home", "San Antonio, TX & Phoenix, AZ",
    "Interior Design Consultant / Studio Coordinator", "Feb 2019 – Oct 2022",
    [
        "Consulted with hundreds of new homeowners to position budgets for cohesive, marketable interior packages that drove upgrade revenue.",
        "Selected as a national trainer to lead design teams through a company-wide digital pivot, developing virtual selection systems during the COVID-19 transition.",
        "Managed the full lifecycle of the build process, from design and budget through change orders and product guide maintenance.",
    ]
)

# Bassett
add_position(
    "Bassett Furniture", "San Antonio, TX",
    "Interior Design Consultant", "Sep 2016 – Aug 2018",
    [
        "Developed a relationship selling framework through in-home consultations, building long-term client loyalty and maximizing per-client revenue.",
        "Led personalized residential design projects, specifying furniture, textiles, and material combinations to deliver collected, cohesive aesthetics.",
    ]
)


# ══════════════════════════════════════════════════════════
#  EARLY CAREER
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Early Career")
set_font(run, name=HEADING_FONT, size=14, color=NAVY)
add_spacing(p, before=6, after=1)
add_thin_line()

# Talent Acquisition
p = doc.add_paragraph()
run = p.add_run("Talent Acquisition Specialist")
set_font(run, size=10, color=DARK_TEXT, bold=True)
run = p.add_run("  |  Investment Professionals, Inc.")
set_font(run, size=9, color=MOCHA)
add_spacing(p, before=3, after=1)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.first_line_indent = Inches(-0.15)
run = p.add_run("●  ")
set_font(run, size=6, color=MOCHA)
run = p.add_run("Developed expertise in networking, behavioral interviewing, and process streamlining for professional recruiting.")
set_font(run, size=11, color=DARK_TEXT)
add_spacing(p, before=0, after=3, line=14)

# EA
p = doc.add_paragraph()
run = p.add_run("Executive Assistant to CEO/COO")
set_font(run, size=10, color=DARK_TEXT, bold=True)
run = p.add_run("  |  Real Estate Training International")
set_font(run, size=9, color=MOCHA)
add_spacing(p, before=3, after=1)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.first_line_indent = Inches(-0.15)
run = p.add_run("●  ")
set_font(run, size=6, color=MOCHA)
run = p.add_run("Supported the scaling of the company from 5 to 400+ employees and closed $9.5M in pending accounts.")
set_font(run, size=11, color=DARK_TEXT)
add_spacing(p, before=0, after=3, line=14)


# ══════════════════════════════════════════════════════════
#  EDUCATION & AFFILIATIONS
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Education & Affiliations")
set_font(run, name=HEADING_FONT, size=14, color=NAVY)
add_spacing(p, before=6, after=1)
add_thin_line()

p = doc.add_paragraph()
run = p.add_run("San Antonio College")
set_font(run, size=10, color=DARK_TEXT, bold=True)
run = p.add_run("  |  Interior Design Coursework")
set_font(run, size=9.5, color=MOCHA)
add_spacing(p, before=3, after=2)

p = doc.add_paragraph()
run = p.add_run("Active Member")
set_font(run, size=10, color=DARK_TEXT, bold=True)
run = p.add_run("  |  San Antonio Construction and Real Estate Networking Groups")
set_font(run, size=9.5, color=MOCHA)
add_spacing(p, before=2, after=2)

# Tools line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TOOLS:  ")
set_font(run, name=BODY_FONT, size=8, color=NAVY, bold=True)
run.font.letter_spacing = Pt(1.5)
run = p.add_run("Canva  ·  Beautiful.ai  ·  SketchUp  ·  Salesforce  ·  MS Office Suite")
set_font(run, size=8.5, color=MOCHA)
add_spacing(p, before=6, after=0)

# Portfolio footer with thin line separator
add_thin_line()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PORTFOLIO")
set_font(run, name=BODY_FONT, size=8, color=NAVY, bold=True)
run.font.letter_spacing = Pt(1.5)
run = p.add_run("    hchcportfolio.vercel.app")
set_font(run, size=9.5, color=MOCHA)
add_spacing(p, before=2, after=0)


# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), "EnaDodski_Resume_2026_v2.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
