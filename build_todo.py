from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

navy = RGBColor(0x1B, 0x2A, 0x4A)
mocha = RGBColor(0xA3, 0x77, 0x62)
espresso = RGBColor(0x3C, 0x2A, 0x21)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = navy
    p.paragraph_format.space_after = Pt(2)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = mocha
    run.font.all_caps = True
    p.paragraph_format.space_after = Pt(20)

def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = navy
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

def add_subsection(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = espresso
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_checkbox(text, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run('\u2610  ' + text)
    run.font.size = Pt(11)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)

def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = RGBColor(0xD4, 0xC5, 0xB2)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# === DOCUMENT CONTENT ===

add_title('HCHC To Do')
add_subtitle('Hill Country Home Concepts  |  Setup Requirements')

add_body('What needs to happen before the full platform can be built. Items are in priority order within each section.')

add_divider()

# --- SECTION 1 ---
add_section_header('1. Immediate — Before Anything Else')

add_subsection('Domain and Hosting')
add_checkbox('Register hillcountryhomeconcepts.com (Namecheap, Google Domains, or Cloudflare)')
add_checkbox('Point custom domain to Vercel (keep free tier for now)')

add_subsection('LLC Formation')
add_checkbox('File LLC in Texas (SOS.Texas.gov, approximately $300 filing fee)')
add_checkbox('Get EIN from IRS (free, immediate online at irs.gov)')
add_checkbox('Open business bank account (Chase, Mercury, or Relay)')

add_subsection('Business Email')
add_checkbox('Set up ena@hillcountryhomeconcepts.com')
add_body('Google Workspace ($7/month) or Zoho (free tier). Replaces Gmail on all business materials.')

add_subsection('Brand Assets')
add_checkbox('Final HCHC logo (wordmark + icon)')
add_checkbox('Favicon (small square version for browser tab)')
add_body('Needed for: website header, presentations, watermarks, business cards, everything.')

add_divider()

# --- SECTION 2 ---
add_section_header('2. Before Builder Portal Goes Live')

add_subsection('Calendar Booking')
add_checkbox('Set up Calendly or Cal.com account (free tier)')
add_checkbox('Create booking link for Design Consultation appointments')
add_body('Will be embedded directly into the website.')

add_subsection('Payment Processing')
add_checkbox('Create Stripe account (free to set up, charges per transaction)')
add_body('Handles subscription billing for the presentation engine and one-time service fees. Requires: business bank account, EIN, LLC.')

add_subsection('Cloud Storage and Authentication')
add_body('No setup needed from Ena. I will configure Vercel Blob or Cloudflare R2 for file storage (Revit uploads, brand books, presentation outputs) and Clerk or Supabase for user login and accounts. Both have free tiers that cover early users.')

add_divider()

# --- SECTION 3 ---
add_section_header('3. Before Presentation Engine MVP')

add_subsection('Material Database')
add_checkbox('Build initial material library spreadsheet')
add_body('Start with 20 to 30 materials per category. Each entry needs:')
add_body('  - Material name')
add_body('  - Category (flooring, tile, cabinet, countertop, hardware, paint, fixture)')
add_body('  - Color tone (warm / cool / neutral)')
add_body('  - Style tags (hill country modern, warm transitional, clean contemporary, classic neutral, elevated builder grade)')
add_body('  - Price tier (standard / upgrade / premium)')
add_body('  - Image (photo of sample or product photo)')
add_body('  - Vendor and SKU (optional)')

add_subsection('Design Rules')
add_checkbox('Document pairing rules (can be conversational, I will structure them)')
add_body('This is the proprietary methodology, the core IP of the business. Needs to cover:')
add_body('  - Which tones pair with which')
add_body('  - Which styles restrict which materials')
add_body('  - What constitutes a conflict (e.g., cool white cabinet + warm honey floor)')
add_body('  - Budget tier constraints')

add_subsection('AI Image Generation (Phase 3, Not Needed for MVP)')
add_checkbox('Research current AI rendering options for Revit/SketchUp files')
add_body('Candidates: Midjourney, DALL-E, Stable Diffusion (local via Forge), or specialized architectural rendering APIs. This comes later.')

add_divider()

# --- COST TABLE ---
add_section_header('Monthly Cost Summary (After Initial Setup)')

table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'

headers = ['Service', 'Cost', 'Notes']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

data = [
    ('Domain', '$12/year', 'One-time annual'),
    ('Vercel hosting', '$0', 'Free tier'),
    ('Google Workspace email', '$7/month', 'Business email'),
    ('Stripe', '2.9% + $0.30/txn', 'Only when charging'),
    ('Calendly', '$0', 'Free tier'),
    ('Auth provider', '$0', 'Free tier'),
    ('File storage', '$1 to $5/month', 'Scales with usage'),
]

for i, (service, cost, notes) in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = service
    row.cells[1].text = cost
    row.cells[2].text = notes
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

add_body('')
add_body('Total fixed cost before any clients: approximately $8/month.')

add_divider()

# --- PRIORITY ORDER ---
add_section_header('Priority Order')
priorities = [
    'Register domain',
    'Form LLC + get EIN',
    'Set up business email',
    'Get logo finalized',
    'Set up Calendly for booking',
    'Set up Stripe (needs LLC + bank account)',
    'Build material spreadsheet (ongoing, feeds presentation engine)',
    'Document design rules (ongoing, feeds presentation engine)',
]
for i, item in enumerate(priorities, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}.  {item}')
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)

add_body('')
add_body('Items 7 and 8 can happen in parallel with everything else. They feed the presentation engine MVP.')

add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Hill Country Home Concepts  |  Confidential  |  April 2026')
run.font.size = Pt(9)
run.font.color.rgb = mocha

doc.save(r'C:\Users\Ajdod\OneDrive\Desktop\Claude\Ena\HCHC To Do.docx')
print('Done.')
